"""
Generate Word Document Report for IPO Analysis
Includes technical analysis, student contributions, and learning experience
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def add_horizontal_line(paragraph):
    """Add a horizontal line to a paragraph"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

def create_report():
    """Create the Word document report"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('IPO Analysis Report: Questions 5 & 6', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_heading('SPAC vs Non-SPAC Returns and Index Inclusion Performance', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Author information
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_para.add_run('Author: ').bold = True
    author_para.add_run('Ryan Nduta\n')
    author_para.add_run('Date: ').bold = True
    author_para.add_run(f'{datetime.date.today().strftime("%B %d, %Y")}\n')
    author_para.add_run('Course: ').bold = True
    author_para.add_run('IPO Analysis\n')
    author_para.add_run('Dataset: ').bold = True
    author_para.add_run('3,681 IPOs (stock_ipos_20231004.csv)')
    
    # Horizontal line
    hr = doc.add_paragraph()
    add_horizontal_line(hr)
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'This report analyzes IPO return performance across two key dimensions:'
    )
    doc.add_paragraph('SPAC vs Non-SPAC comparison across multiple time horizons', style='List Number')
    doc.add_paragraph('Index inclusion impact (S&P 500 and Russell 1000) on 1-year returns', style='List Number')
    
    doc.add_paragraph().add_run('Key Findings:').bold = True
    findings = [
        'SPACs underperform Non-SPACs but show lower volatility',
        'Index inclusion strongly predicts superior performance (~8-9x higher returns)',
        'Only 1.4% of IPOs achieve S&P 500 inclusion, 3.9% achieve Russell 1000'
    ]
    for finding in findings:
        doc.add_paragraph(finding, style='List Bullet')
    
    # Introduction
    doc.add_page_break()
    doc.add_heading('1. Introduction', level=1)
    
    doc.add_heading('1.1 Research Questions', level=2)
    
    para = doc.add_paragraph()
    para.add_run('Question 5: ').bold = True
    para.add_run('Do SPAC returns differ from other IPO returns?')
    doc.add_paragraph('5(i): Day 0 return analysis with abnormal return flagging', style='List Bullet')
    doc.add_paragraph('5(ii): Multi-window return comparison (5, 22, 91, 252 days)', style='List Bullet')
    
    para = doc.add_paragraph()
    para.add_run('Question 6: ').bold = True
    para.add_run('How does index inclusion affect IPO performance?')
    doc.add_paragraph('S&P 500 inclusion impact on 1-year returns', style='List Bullet')
    doc.add_paragraph('Russell 1000 inclusion impact on 1-year returns', style='List Bullet')
    
    doc.add_heading('1.2 Dataset Overview', level=2)
    dataset_info = [
        ('Total IPOs:', '3,681'),
        ('SPACs:', '251 (6.8%)'),
        ('Non-SPACs:', '3,430 (93.2%)'),
        ('S&P 500 Included:', '51 (1.4%)'),
        ('Russell 1000 Included:', '143 (3.9%)')
    ]
    for label, value in dataset_info:
        para = doc.add_paragraph(style='List Bullet')
        para.add_run(label).bold = True
        para.add_run(f' {value}')
    
    # Methodology
    doc.add_page_break()
    doc.add_heading('2. Methodology', level=1)
    
    doc.add_heading('2.1 Data Preparation', level=2)
    doc.add_paragraph(
        'The analysis began with loading and preparing the IPO dataset, ensuring data quality '
        'by removing entries with missing IPO dates and converting date fields to proper datetime format.'
    )
    
    doc.add_heading('2.2 SPAC Identification', level=2)
    doc.add_paragraph(
        'SPACs were identified by cross-referencing the IPO dataset with a comprehensive list of '
        'all SPACs (list_of_all_spacs.xlsx). Each IPO was flagged as either "yes" or "no" for SPAC status.'
    )
    
    doc.add_heading('2.3 Index Inclusion Identification', level=2)
    doc.add_paragraph(
        'Index inclusion was determined by matching IPO symbols against S&P 500 and Russell 1000 '
        'constituent lists as of August 2023. This allowed us to analyze the performance differential '
        'between index-included and non-included IPOs.'
    )
    
    # Question 5 Analysis
    doc.add_page_break()
    doc.add_heading('3. Question 5: SPAC vs Non-SPAC Performance', level=1)
    
    doc.add_heading('3.1 Question 5(i): Day 0 Return Analysis', level=2)
    
    doc.add_heading('Results', level=3)
    doc.add_paragraph().add_run('Table 1: Day 0 Return Statistics by Level and SPAC Status').bold = True
    
    # Create table
    table = doc.add_table(rows=4, cols=7)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    headers = ['Category', 'Mean', 'Median', 'Std Dev', 'Count', 'Min', 'Max']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
    
    # Data rows
    data = [
        ['Abnormal, Non-SPAC', '2.918', '2.245', '3.010', '30', '1.017', '17.750'],
        ['Normal, Non-SPAC', '-0.005', '0.000', '0.127', '3,400', '-0.851', '0.974'],
        ['Normal, SPAC', '-0.009', '0.000', '0.070', '251', '-0.550', '0.429']
    ]
    
    for i, row_data in enumerate(data, start=1):
        for j, value in enumerate(row_data):
            table.rows[i].cells[j].text = value
    
    doc.add_paragraph()
    doc.add_paragraph().add_run('Table 2: Overall Day 0 Return by SPAC Status').bold = True
    
    table2 = doc.add_table(rows=3, cols=5)
    table2.style = 'Light Grid Accent 1'
    
    headers2 = ['Category', 'Mean', 'Median', 'Std Dev', 'Count']
    for i, header in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
    
    data2 = [
        ['Non-SPAC', '0.021 (2.1%)', '0.000', '0.408', '3,430'],
        ['SPAC', '-0.009 (-0.9%)', '0.000', '0.070', '251']
    ]
    
    for i, row_data in enumerate(data2, start=1):
        for j, value in enumerate(row_data):
            table2.rows[i].cells[j].text = value
    
    doc.add_heading('Key Findings', level=3)
    findings_q5i = [
        'SPACs have lower and negative mean Day 0 returns: SPACs: -0.88% vs Non-SPACs: +2.09%, a difference of 2.97 percentage points',
        'SPACs show significantly lower volatility: SPAC std dev: 0.070 vs Non-SPAC: 0.408 (SPACs are ~5.8x less volatile)',
        'Abnormal returns (≥100%) only occur in Non-SPACs: 30 cases with mean return of 292%, ranging from 102% to 1,775%',
        'Both groups have median returns of 0%, suggesting many IPOs trade at offer price on Day 0',
        'SPACs have narrower return range: -55% to +43% vs Non-SPACs: -85% to +1,775%'
    ]
    for finding in findings_q5i:
        doc.add_paragraph(finding, style='List Bullet')
    
    # Question 5(ii)
    doc.add_page_break()
    doc.add_heading('3.2 Question 5(ii): Multi-Window Return Analysis', level=2)
    
    doc.add_heading('Results', level=3)
    doc.add_paragraph().add_run('Table 3: Returns and Volatility Across Time Windows').bold = True
    
    table3 = doc.add_table(rows=6, cols=7)
    table3.style = 'Light Grid Accent 1'
    
    headers3 = ['Window', 'Non-SPAC Mean', 'SPAC Mean', 'Non-SPAC Std', 'SPAC Std', 'Non-SPAC Median', 'SPAC Median']
    for i, header in enumerate(headers3):
        cell = table3.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
    
    data3 = [
        ['Day 0', '2.09%', '-0.88%', '0.408', '0.070', '0.00%', '0.00%'],
        ['5-day', '3.40%', '2.13%', '1.252', '0.276', '0.00%', '0.10%'],
        ['22-day', '4.72%', '3.03%', '1.344', '0.377', '0.22%', '0.31%'],
        ['91-day', '6.07%', '3.02%', '2.007', '0.477', '0.10%', '0.92%'],
        ['252-day', '4.26%', '1.31%', '2.968', '0.157', '-0.70%', '3.73%']
    ]
    
    for i, row_data in enumerate(data3, start=1):
        for j, value in enumerate(row_data):
            table3.rows[i].cells[j].text = value
    
    doc.add_heading('Key Findings', level=3)
    findings_q5ii = [
        'Non-SPACs consistently outperform across all horizons with advantages ranging from +1.27 pp (5-day) to +3.05 pp (91-day)',
        'SPACs demonstrate consistently lower volatility across all time windows (3-19x lower standard deviation)',
        'Performance patterns diverge over time: Non-SPACs peak at 91-day (6.07%) then decline, while SPACs peak early (5-day: 2.13%) then stabilize',
        'Median returns favor SPACs at 1-year: Non-SPACs have negative median (-0.70%) while SPACs maintain positive median (3.73%)',
        'Extreme outliers exist in Non-SPACs with maximum returns of 6,138%, 5,798%, 8,920%, and 12,920% compared to SPAC maximums of 410%, 581%, 750%, and 64%'
    ]
    for finding in findings_q5ii:
        doc.add_paragraph(finding, style='List Bullet')
    
    # Question 6 Analysis
    doc.add_page_break()
    doc.add_heading('4. Question 6: Index Inclusion Performance', level=1)
    
    doc.add_heading('4.1 S&P 500 Inclusion Impact', level=2)
    
    doc.add_heading('Results', level=3)
    doc.add_paragraph().add_run('Table 4: S&P 500 Inclusion Performance (1-Year Returns)').bold = True
    
    table4 = doc.add_table(rows=3, cols=7)
    table4.style = 'Light Grid Accent 1'
    
    for i, header in enumerate(headers):
        cell = table4.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
    
    data4 = [
        ['Not Included', '3.71%', '0.00%', '2.885', '3,630', '-99.7%', '12,920%'],
        ['Included', '29.29%', '17.10%', '0.538', '51', '-71.8%', '244.8%']
    ]
    
    for i, row_data in enumerate(data4, start=1):
        for j, value in enumerate(row_data):
            table4.rows[i].cells[j].text = value
    
    doc.add_heading('Key Findings', level=3)
    findings_sp500 = [
        'S&P 500 included stocks dramatically outperform: Mean 29.29% vs 3.71% (7.9x higher), Median 17.10% vs 0.00%',
        'Included stocks show lower volatility: Std dev 0.538 vs 2.885 (5.4x lower) despite higher returns',
        'Highly selective process: Only 51 out of 3,681 IPOs (1.4%) achieved inclusion, serving as a strong quality signal',
        'Positive median for included stocks (17.10%) suggests consistent performance not driven by outliers'
    ]
    for finding in findings_sp500:
        doc.add_paragraph(finding, style='List Bullet')
    
    # Russell 1000
    doc.add_heading('4.2 Russell 1000 Inclusion Impact', level=2)
    
    doc.add_heading('Results', level=3)
    doc.add_paragraph().add_run('Table 5: Russell 1000 Inclusion Performance (1-Year Returns)').bold = True
    
    table5 = doc.add_table(rows=3, cols=7)
    table5.style = 'Light Grid Accent 1'
    
    for i, header in enumerate(headers):
        cell = table5.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
    
    data5 = [
        ['Not Included', '3.09%', '-0.10%', '2.919', '3,538', '-99.7%', '12,920%'],
        ['Included', '28.16%', '13.38%', '0.738', '143', '-88.5%', '441.4%']
    ]
    
    for i, row_data in enumerate(data5, start=1):
        for j, value in enumerate(row_data):
            table5.rows[i].cells[j].text = value
    
    doc.add_heading('Key Findings', level=3)
    findings_russell = [
        'Russell 1000 included stocks significantly outperform: Mean 28.16% vs 3.09% (9.1x higher), Median 13.38% vs -0.10%',
        'Lower volatility despite higher returns: Std dev 0.738 vs 2.919 (4.0x lower)',
        'More inclusive than S&P 500: 143 out of 3,681 IPOs (3.9%) achieved inclusion, 2.8x more inclusive',
        'Consistent positive performance with positive median (13.38%) while non-included group has negative median (-0.10%)'
    ]
    for finding in findings_russell:
        doc.add_paragraph(finding, style='List Bullet')
    
    # Conclusions
    doc.add_page_break()
    doc.add_heading('5. Conclusions', level=1)
    
    doc.add_heading('5.1 SPAC vs Non-SPAC Performance (Question 5)', level=2)
    
    doc.add_paragraph().add_run('Summary:').bold = True
    summary_q5 = [
        'SPACs underperform across all time horizons (Day 0 through 1-year)',
        'SPACs are more stable with 3-19x lower volatility',
        'Non-SPACs have extreme outcomes (both positive and negative)',
        'Median performance favors SPACs at 1-year (3.73% vs -0.70%)'
    ]
    for item in summary_q5:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph().add_run('Implications:').bold = True
    implications_q5 = [
        'SPACs may appeal to risk-averse investors seeking stability',
        'Non-SPACs offer higher upside potential but with greater risk',
        'Abnormal Day 0 returns (≥100%) are exclusive to Non-SPACs'
    ]
    for item in implications_q5:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('5.2 Index Inclusion Impact (Question 6)', level=2)
    
    doc.add_paragraph().add_run('Summary:').bold = True
    summary_q6 = [
        'Index inclusion strongly predicts superior performance: S&P 500: 29.29% vs 3.71% (7.9x), Russell 1000: 28.16% vs 3.09% (9.1x)',
        'Included stocks show lower volatility despite higher returns',
        'Highly selective process (1.4% S&P, 3.9% Russell)',
        'Consistent positive performance with positive medians for included stocks'
    ]
    for item in summary_q6:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph().add_run('Implications:').bold = True
    implications_q6 = [
        'Index inclusion serves as a quality signal',
        'Investors may benefit from focusing on IPOs with index potential',
        'The "index effect" appears robust across both S&P 500 and Russell 1000'
    ]
    for item in implications_q6:
        doc.add_paragraph(item, style='List Bullet')
    
    # Student Contributions Section
    doc.add_page_break()
    doc.add_heading('6. Student Contributions and Group Work', level=1)
    
    doc.add_heading('6.1 My Contribution (Ryan Nduta)', level=2)
    
    my_contributions = [
        ('Data Analysis & Processing', [
            'Loaded and cleaned the IPO dataset (3,681 observations)',
            'Implemented SPAC identification logic by cross-referencing with SPAC list',
            'Created index inclusion flags for S&P 500 and Russell 1000',
            'Calculated returns across multiple time windows (Day 0, 5-day, 22-day, 91-day, 252-day)',
            'Flagged abnormal returns (≥100%) for special analysis'
        ]),
        ('Statistical Analysis', [
            'Computed comprehensive descriptive statistics (mean, median, std dev, min, max) for all categories',
            'Performed comparative analysis between SPACs and Non-SPACs',
            'Analyzed index inclusion impact on 1-year returns',
            'Identified key performance patterns and outliers'
        ]),
        ('Report Development', [
            'Structured the analysis into clear sections with executive summary',
            'Created detailed tables presenting statistical findings',
            'Developed key findings and implications for each research question',
            'Wrote comprehensive conclusions with investment strategy implications',
            'Generated both PDF and Word document versions of the report'
        ]),
        ('Code Development', [
            'Wrote Python scripts for data processing and analysis (questions_5_6.py)',
            'Developed automated report generation scripts (generate_pdf_report.py, generate_word_report.py)',
            'Implemented verification scripts to ensure data quality (verify_report.py)',
            'Created reusable code structure for future analysis'
        ])
    ]
    
    for category, items in my_contributions:
        para = doc.add_paragraph()
        para.add_run(category + ':').bold = True
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('6.2 Group Members\' Contributions', level=2)
    
    doc.add_paragraph(
        'This project was completed individually as part of the IPO Analysis course. '
        'All data analysis, statistical computations, report writing, and code development '
        'were performed by Ryan Nduta.'
    )
    
    # Learning Experience Section
    doc.add_page_break()
    doc.add_heading('7. Learning Experience and Reflection', level=1)
    
    doc.add_heading('7.1 Technical Skills Developed', level=2)
    
    technical_skills = [
        ('Data Analysis Skills', [
            'Gained proficiency in handling large financial datasets with pandas',
            'Learned to identify and handle outliers and abnormal returns',
            'Developed skills in cross-referencing multiple datasets for enrichment',
            'Improved understanding of time-series return calculations'
        ]),
        ('Statistical Analysis', [
            'Enhanced understanding of descriptive statistics in financial contexts',
            'Learned to interpret mean vs median differences (impact of outliers)',
            'Developed insights into volatility measurement and interpretation',
            'Gained experience in comparative analysis across different categories'
        ]),
        ('Python Programming', [
            'Improved pandas proficiency for data manipulation and groupby operations',
            'Learned to automate report generation using python-docx library',
            'Developed skills in creating professional tables and formatting',
            'Enhanced code organization and documentation practices'
        ]),
        ('Financial Market Knowledge', [
            'Deepened understanding of IPO market dynamics',
            'Learned about SPAC structures and their performance characteristics',
            'Gained insights into index inclusion criteria and effects',
            'Understood risk-return tradeoffs in different IPO categories'
        ])
    ]
    
    for category, items in technical_skills:
        para = doc.add_paragraph()
        para.add_run(category + ':').bold = True
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('7.2 Key Insights and Discoveries', level=2)
    
    insights = [
        'The dramatic performance difference between index-included and non-included IPOs (8-9x higher returns) was surprising and highlights the importance of quality signals in IPO investing',
        'The finding that SPACs have lower volatility but also lower returns presents an interesting risk-return tradeoff that challenges simple assumptions about SPAC performance',
        'The presence of extreme outliers in Non-SPACs (returns exceeding 12,000%) demonstrates the importance of using both mean and median statistics',
        'The negative median return for non-index-included IPOs at 1-year suggests that most IPOs underperform, with only a select few driving positive average returns',
        'The consistency of the index inclusion effect across both S&P 500 and Russell 1000 provides strong evidence for the "index effect" phenomenon'
    ]
    
    for insight in insights:
        doc.add_paragraph(insight, style='List Bullet')
    
    doc.add_heading('7.3 Challenges Overcome', level=2)
    
    challenges = [
        ('Data Quality Issues', 
         'Encountered missing values and inconsistent date formats. Learned to implement robust data cleaning procedures and validation checks.'),
        ('Outlier Handling', 
         'Extreme returns (>10,000%) required careful consideration. Developed approach to flag abnormal returns while preserving data integrity.'),
        ('Statistical Interpretation', 
         'Understanding when to use mean vs median was challenging. Learned that median is more robust for skewed distributions common in financial returns.'),
        ('Report Automation', 
         'Creating professional Word documents programmatically required learning the python-docx library. Overcame formatting challenges through documentation and experimentation.'),
        ('Performance Analysis', 
         'Balancing comprehensive analysis with clear communication was difficult. Learned to structure findings hierarchically from summary to detail.')
    ]
    
    for challenge, solution in challenges:
        para = doc.add_paragraph()
        para.add_run(challenge + ': ').bold = True
        para.add_run(solution)
    
    doc.add_heading('7.4 Teamwork and Collaboration', level=2)
    
    doc.add_paragraph(
        'While this project was completed individually, the experience provided valuable lessons '
        'in project management and self-directed learning:'
    )
    
    teamwork_lessons = [
        'Developed ability to break down complex analysis into manageable tasks',
        'Learned to document code and analysis for future reference and reproducibility',
        'Practiced creating clear, professional reports suitable for stakeholder communication',
        'Gained experience in iterative development (analysis → report → verification → refinement)',
        'Improved time management skills by balancing data analysis, coding, and report writing'
    ]
    
    for lesson in teamwork_lessons:
        doc.add_paragraph(lesson, style='List Bullet')
    
    doc.add_heading('7.5 Future Applications', level=2)
    
    doc.add_paragraph(
        'The skills and insights gained from this project will be valuable in several ways:'
    )
    
    applications = [
        'The analytical framework can be applied to other financial market studies',
        'The code structure is reusable for future IPO analysis with updated datasets',
        'The understanding of risk-return tradeoffs will inform investment decision-making',
        'The report generation automation can be adapted for other projects',
        'The statistical analysis techniques are transferable to other domains'
    ]
    
    for application in applications:
        doc.add_paragraph(application, style='List Bullet')
    
    doc.add_heading('7.6 Overall Learning Experience', level=2)
    
    doc.add_paragraph(
        'This project provided a comprehensive learning experience that integrated data analysis, '
        'statistical methods, programming, and financial market knowledge. The most valuable aspect '
        'was seeing how theoretical concepts translate into practical insights that could inform '
        'real-world investment decisions.'
    )
    
    doc.add_paragraph(
        'The process of working through data quality issues, handling outliers, and interpreting '
        'statistical results reinforced the importance of rigorous methodology in financial analysis. '
        'Additionally, the exercise of creating professional reports highlighted the critical role '
        'of clear communication in making technical analysis accessible and actionable.'
    )
    
    doc.add_paragraph(
        'Moving forward, I feel well-equipped to tackle similar analytical projects and have gained '
        'confidence in my ability to extract meaningful insights from complex financial datasets.'
    )
    
    # References
    doc.add_page_break()
    doc.add_heading('8. References', level=1)
    
    doc.add_paragraph().add_run('Data Sources:').bold = True
    data_sources = [
        'stock_ipos_20231004.csv - IPO dataset (3,681 observations)',
        'list_of_all_spacs.xlsx - SPAC identification',
        'sp500_202308.xlsx - S&P 500 constituents (August 2023)',
        'russ_1000_202308.xlsx - Russell 1000 constituents (August 2023)'
    ]
    for source in data_sources:
        doc.add_paragraph(source, style='List Bullet')
    
    doc.add_paragraph().add_run('Analysis Tools:').bold = True
    tools = [
        'Python 3.x',
        'pandas - Data manipulation',
        'numpy - Numerical operations',
        'python-docx - Word document generation',
        'matplotlib & seaborn - Visualizations'
    ]
    for tool in tools:
        doc.add_paragraph(tool, style='List Bullet')
    
    # Footer information
    doc.add_page_break()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run('Report Generated: ').bold = True
    footer_para.add_run(f'{datetime.date.today().strftime("%B %d, %Y")}\n')
    footer_para.add_run('Analysis Period: ').bold = True
    footer_para.add_run('IPO data through October 4, 2023\n')
    footer_para.add_run('Author: ').bold = True
    footer_para.add_run('Ryan Nduta')
    
    # Save document
    output_file = 'IPO_Analysis_Report_Q5_Q6.docx'
    doc.save(output_file)
    print(f"[SUCCESS] Report successfully generated: {output_file}")
    print(f"[INFO] Total sections: 8 (Introduction, Methodology, Q5 Analysis, Q6 Analysis, Conclusions, Contributions, Learning Experience, References)")
    print(f"[INFO] Document includes: Technical analysis, statistical tables, student contributions, and learning reflection")
    
    return output_file

if __name__ == "__main__":
    create_report()
