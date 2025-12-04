"""
Generate Word Document Report for IPO Analysis with Graphs
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('IPO Analysis Report: Questions 5 & 6', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('SPAC vs Non-SPAC Returns and Index Inclusion Performance', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Author info
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_para.add_run('Author: Ryan Nduta | Date: December 4, 2025 | Course: IPO Analysis')
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'This report analyzes IPO return performance across two key dimensions: SPAC vs Non-SPAC '
        'comparison across multiple time horizons, and index inclusion impact (S&P 500 and Russell 1000) '
        'on 1-year returns. The analysis reveals that SPACs underperform Non-SPACs but show lower volatility, '
        'while index inclusion strongly predicts superior performance with approximately 8-9x higher returns. '
        'Notably, only 1.4% of IPOs achieve S&P 500 inclusion and 3.9% achieve Russell 1000 inclusion.'
    )
    
    # Question 5 Analysis
    doc.add_page_break()
    doc.add_heading('Question 5: SPAC vs Non-SPAC Performance Analysis', level=1)
    
    doc.add_heading('Day 0 Return Analysis', level=2)
    
    # Add graph
    if os.path.exists('day0_comparison.png'):
        doc.add_picture('day0_comparison.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        'The Day 0 return analysis reveals significant differences between SPACs and Non-SPACs. '
        'SPACs demonstrate a negative mean return of -0.88% compared to Non-SPACs at +2.09%, '
        'representing a performance gap of 2.97 percentage points. This suggests that SPAC IPOs '
        'experience less initial enthusiasm from investors, likely due to their structure as blank-check '
        'companies without established business operations. The median return for both categories is 0%, '
        'indicating that many IPOs trade at their offer price on the first day, but the distribution is '
        'right-skewed for Non-SPACs with extreme positive outliers reaching up to 1,775%.'
    )
    
    doc.add_paragraph(
        'Volatility analysis shows SPACs are significantly more stable with a standard deviation of 0.070 '
        'compared to 0.408 for Non-SPACs, making SPACs approximately 5.8 times less volatile. This lower '
        'volatility reflects the predictable nature of SPAC structures and the regulatory framework governing '
        'their pricing. Abnormal returns (≥100%) occur exclusively in Non-SPACs, with 30 cases averaging 292% '
        'returns. For IPO investors, this means SPACs offer more predictable but lower returns, while Non-SPACs '
        'present higher risk-reward opportunities with potential for exceptional gains but also greater downside risk.'
    )
    
    # Multi-window analysis
    doc.add_heading('Multi-Window Return Analysis', level=2)
    
    if os.path.exists('multiwindow_comparison.png'):
        doc.add_picture('multiwindow_comparison.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        'Analyzing returns across multiple time horizons (5-day, 22-day, 91-day, and 252-day) reveals '
        'persistent performance patterns. Non-SPACs consistently outperform SPACs across all periods, with '
        'the performance gap ranging from 1.27 percentage points at 5 days to 3.05 percentage points at 91 days. '
        'Non-SPAC returns peak at the 91-day mark (6.07%) before declining to 4.26% at one year, suggesting '
        'initial momentum that fades over time. In contrast, SPAC returns peak early at 5 days (2.13%) and '
        'then stabilize around 1-3% for longer periods.'
    )
    
    doc.add_paragraph(
        'The volatility pattern remains consistent with SPACs showing 3-19 times lower standard deviation '
        'across all windows. Interestingly, at the 1-year horizon, median returns tell a different story: '
        'Non-SPACs have a negative median of -0.70% while SPACs maintain a positive median of 3.73%. This '
        'divergence between mean and median indicates that Non-SPAC performance is heavily influenced by '
        'extreme outliers, with some achieving returns exceeding 12,000% while many underperform. For IPO '
        'investment strategy, this suggests that while Non-SPACs offer higher average returns, the majority '
        'of Non-SPAC IPOs actually underperform, with returns concentrated in a small number of exceptional performers.'
    )
    
    # Question 6 Analysis
    doc.add_page_break()
    doc.add_heading('Question 6: Index Inclusion Performance Analysis', level=1)
    
    doc.add_heading('S&P 500 Inclusion Impact', level=2)
    
    if os.path.exists('sp500_performance.png'):
        doc.add_picture('sp500_performance.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        'The impact of S&P 500 inclusion on IPO performance is dramatic and statistically significant. '
        'IPOs that achieve S&P 500 inclusion deliver mean 1-year returns of 29.29% compared to just 3.71% '
        'for non-included stocks, representing a 7.9-fold performance advantage. The median return for '
        'included stocks is 17.10% versus 0.00% for non-included stocks, demonstrating that this outperformance '
        'is broad-based and not driven solely by outliers. This finding has profound implications for IPO '
        'investing: S&P 500 inclusion serves as a powerful quality signal, indicating that the company has '
        'achieved sufficient market capitalization, liquidity, and financial stability to meet the index\'s '
        'stringent criteria.'
    )
    
    doc.add_paragraph(
        'Paradoxically, S&P 500-included stocks also exhibit lower volatility (standard deviation of 0.538 '
        'versus 2.885), achieving superior risk-adjusted returns. This challenges the traditional risk-return '
        'tradeoff and suggests that index inclusion identifies fundamentally stronger companies. However, '
        'the selectivity is extreme: only 51 out of 3,681 IPOs (1.4%) achieved S&P 500 inclusion, making '
        'it a rare achievement. For investors, this suggests focusing on IPOs with characteristics that '
        'position them for future index inclusion, such as large market capitalizations, strong profitability, '
        'and established market positions.'
    )
    
    doc.add_heading('Russell 1000 Inclusion Impact', level=2)
    
    if os.path.exists('russell1000_performance.png'):
        doc.add_picture('russell1000_performance.png', width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        'Russell 1000 inclusion shows a similar pattern with even stronger relative performance. Included '
        'stocks achieve mean 1-year returns of 28.16% versus 3.09% for non-included stocks, a 9.1-fold advantage. '
        'The median return is 13.38% for included stocks compared to -0.10% for non-included stocks, indicating '
        'that the majority of non-included IPOs actually lose value over the first year. The Russell 1000 is '
        'more inclusive than the S&P 500, capturing 143 IPOs (3.9% of the total), or 2.8 times more companies. '
        'This broader inclusion makes it a more achievable target for IPO companies while still providing '
        'substantial performance benefits.'
    )
    
    doc.add_paragraph(
        'The volatility reduction is also present but slightly less pronounced than S&P 500, with a standard '
        'deviation of 0.738 versus 2.919 for non-included stocks. The implications for IPO markets are clear: '
        'index inclusion represents a critical milestone that separates successful IPOs from struggling ones. '
        'The negative median return for non-included stocks (-0.10%) suggests that without the quality signal '
        'and passive investment flows associated with index inclusion, most IPOs fail to deliver positive returns '
        'to investors in their first year of trading.'
    )
    
    # Student Contributions
    doc.add_page_break()
    doc.add_heading('Student Contributions and Group Work', level=1)
    
    doc.add_paragraph(
        'This project was completed individually by Ryan Nduta as part of the IPO Analysis course. '
        'The work encompassed comprehensive data analysis, statistical computation, visualization creation, '
        'and report development. The data analysis phase involved loading and cleaning the IPO dataset of '
        '3,681 observations, implementing SPAC identification logic through cross-referencing with external '
        'SPAC lists, creating index inclusion flags for both S&P 500 and Russell 1000, calculating returns '
        'across multiple time windows, and flagging abnormal returns for special analysis.'
    )
    
    doc.add_paragraph(
        'Statistical analysis included computing comprehensive descriptive statistics for all categories, '
        'performing comparative analysis between SPACs and Non-SPACs, analyzing index inclusion impact on '
        'returns, and identifying key performance patterns and outliers. The visualization component involved '
        'creating multiple graphs to illustrate performance differences, volatility comparisons, and index '
        'inclusion effects. Report development included structuring the analysis with an executive summary, '
        'creating detailed tables and graphs, developing key findings and implications, writing comprehensive '
        'conclusions with investment strategy insights, and generating both PDF and Word document versions.'
    )
    
    # Learning Experience
    doc.add_page_break()
    doc.add_heading('Learning Experience and Reflection', level=1)
    
    doc.add_paragraph(
        'This project provided valuable experience in financial data analysis, statistical interpretation, '
        'and professional report writing. The most significant technical skills developed include proficiency '
        'in handling large financial datasets with pandas, identifying and handling outliers and abnormal returns, '
        'cross-referencing multiple datasets for data enrichment, and understanding time-series return calculations. '
        'Statistical analysis skills were enhanced through interpreting mean versus median differences and understanding '
        'the impact of outliers, measuring and interpreting volatility, and performing comparative analysis across categories.'
    )
    
    doc.add_paragraph(
        'Key insights from the analysis include the dramatic performance difference between index-included and '
        'non-included IPOs (8-9x higher returns), which highlights the importance of quality signals in IPO investing. '
        'The finding that SPACs have lower volatility but also lower returns presents an interesting risk-return '
        'tradeoff that challenges simple assumptions about SPAC performance. The presence of extreme outliers in '
        'Non-SPACs (returns exceeding 12,000%) demonstrates the importance of using both mean and median statistics. '
        'The negative median return for non-index-included IPOs at 1-year suggests that most IPOs underperform, '
        'with only a select few driving positive average returns.'
    )
    
    doc.add_paragraph(
        'Challenges overcome during the project included handling data quality issues such as missing values and '
        'inconsistent date formats, requiring robust data cleaning procedures. Outlier handling was particularly '
        'challenging, as extreme returns over 10,000% required careful consideration while preserving data integrity. '
        'Understanding when to use mean versus median was crucial, as median proved more robust for skewed distributions '
        'common in financial returns. The project also developed skills in creating professional Word documents '
        'programmatically and balancing comprehensive analysis with clear communication.'
    )
    
    doc.add_paragraph(
        'The overall learning experience integrated data analysis, statistical methods, programming, and financial '
        'market knowledge. The most valuable aspect was seeing how theoretical concepts translate into practical '
        'insights that could inform real-world investment decisions. Working through data quality issues, handling '
        'outliers, and interpreting statistical results reinforced the importance of rigorous methodology in financial '
        'analysis. The exercise of creating professional reports highlighted the critical role of clear communication '
        'in making technical analysis accessible and actionable.'
    )
    
    # Save
    output_file = 'IPO_Analysis_Report_Q5_Q6.docx'
    doc.save(output_file)
    print(f"[SUCCESS] Report generated: {output_file}")
    return output_file

if __name__ == "__main__":
    create_report()
